from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from dataclasses import dataclass
from typing import List, Optional
from datetime import datetime
from docx2pdf import convert
import os

@dataclass
class MediaAdvisory:
    headline: str
    summary: str
    event_date: str
    event_time: str
    location: str
    contact_name: str
    contact_phone: str
    contact_email: str
    additional_speakers: List[str] = None

@dataclass
class Speaker:
    name: str
    title: str
    organization: str = ""
    speaking_time: str = ""
    notes: Optional[str] = None

@dataclass
class RunOfShow:
    event_title: str
    event_date: str
    event_location: str
    speakers: List[Speaker]

class PressDocGenerator:
    def __init__(self):
        self.doc = None

    def create_media_advisory(self, data: MediaAdvisory) -> str:
        self.doc = Document()
        
        # Set up margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Header - "MEDIA ADVISORY"
        header = self.doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header.add_run("*** MEDIA ADVISORY ***")
        header_run.bold = True
        header_run.font.size = Pt(14)
        
        # Date
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.add_run(f"For {data.event_date}")
        self.doc.add_paragraph()  # Blank line
        
        # Headline
        headline = self.doc.add_paragraph()
        headline.alignment = WD_ALIGN_PARAGRAPH.LEFT
        headline_run = headline.add_run(data.headline.upper())
        headline_run.bold = True
        self.doc.add_paragraph()  # Blank line
        
        # Location and Summary
        location_para = self.doc.add_paragraph()
        location_para.add_run("NEW HAVEN, CT -- ").bold = True
        location_para.add_run(data.summary)
        self.doc.add_paragraph()  # Blank line
        
        # Event Details
        self.doc.add_paragraph("EVENT DETAILS:", style='Normal').bold = True
        self.doc.add_paragraph()  # Blank line
        
        what_para = self.doc.add_paragraph()
        what_para.add_run("WHAT: ").bold = True
        what_para.add_run(data.headline)
        
        when_para = self.doc.add_paragraph()
        when_para.add_run("WHEN: ").bold = True
        when_para.add_run(f"{data.event_date}, {data.event_time}")
        
        where_para = self.doc.add_paragraph()
        where_para.add_run("WHERE: ").bold = True
        where_para.add_run(data.location)
        
        # Speakers section
        if data.additional_speakers:
            who_para = self.doc.add_paragraph()
            who_para.add_run("WHO:").bold = True
            self.doc.add_paragraph("Speakers:").underline = True
            
            for speaker in data.additional_speakers:
                self.doc.add_paragraph(f"• {speaker}", style='List Bullet')
        
        self.doc.add_paragraph()  # Blank line
        self.doc.add_paragraph("###")
        
        # Contact Information
        contact_header = self.doc.add_paragraph()
        contact_header.add_run("Press Contact:").underline = True
        contact = self.doc.add_paragraph()
        contact.add_run(f"{data.contact_name}\n")
        contact.add_run(f"Phone: {data.contact_phone}\n")
        contact.add_run(f"Email: {data.contact_email}")
        
        # Save as DOCX
        docx_path = "temp_advisory.docx"
        self.doc.save(docx_path)
        
        # Convert to PDF
        pdf_path = "media_advisory.pdf"
        convert(docx_path, pdf_path)
        
        # Clean up temp file
        os.remove(docx_path)
        
        return pdf_path

    def create_run_of_show(self, data: RunOfShow) -> str:
        self.doc = Document()
        
        # Set up margins
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title and Event Info
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(data.event_title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(data.event_date).italic = True
        
        location_para = self.doc.add_paragraph()
        location_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        location_para.add_run(data.event_location)
        
        self.doc.add_paragraph()  # Blank line
        
        # Run of Show Header
        ros_header = self.doc.add_paragraph()
        ros_header.add_run("Run of Show (in order of appearance):").bold = True
        ros_header.italic = True
        
        self.doc.add_paragraph("Podium & Speaker will be provided").italic = True
        self.doc.add_paragraph()  # Blank line
        
        # Speakers
        for speaker in data.speakers:
            speaker_para = self.doc.add_paragraph(style='List Bullet')
            name_run = speaker_para.add_run(speaker.name)
            name_run.bold = True
            
            if speaker.title or speaker.organization:
                title_org = f", {speaker.title}"
                if speaker.organization:
                    title_org += f" of {speaker.organization}"
                speaker_para.add_run(title_org)
            
            if speaker.speaking_time:
                speaker_para.add_run(f" ({speaker.speaking_time})")
            
            if speaker.notes:
                notes_para = self.doc.add_paragraph()
                notes_para.add_run(f"Notes: {speaker.notes}").italic = True
        
        # Save as DOCX
        docx_path = "temp_runofshow.docx"
        self.doc.save(docx_path)
        
        # Convert to PDF
        pdf_path = "run_of_show.pdf"
        convert(docx_path, pdf_path)
        
        # Clean up temp file
        os.remove(docx_path)
        
        return pdf_path

# Flask backend to handle document generation
from flask import Flask, request, send_file
import json

app = Flask(__name__)

@app.route('/generate-document', methods=['POST'])
def generate_document():
    data = request.json
    generator = PressDocGenerator()
    
    if data['type'] == 'advisory':
        advisory_data = MediaAdvisory(**data['content'])
        pdf_path = generator.create_media_advisory(advisory_data)
    else:
        show_data = RunOfShow(**data['content'])
        pdf_path = generator.create_run_of_show(show_data)
    
    return send_file(pdf_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
